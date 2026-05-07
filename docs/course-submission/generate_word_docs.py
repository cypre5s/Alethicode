from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).parent
TECH_DOCS = [
    "01-project-plan.md",
    "02-software-requirements-specification.md",
    "03-high-level-design.md",
    "04-detailed-design.md",
    "05-test-plan-and-report.md",
    "06-project-summary.md",
]


def set_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    for name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(30, 39, 97)


def add_cover(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(30, 39, 97)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(14)

    doc.add_paragraph()
    for line in [
        "项目名称：Alethicode 智能在线评测平台",
        "项目定位：面向非计算机专业 Python 初学者的 AI 编程学习系统",
        "文档版本：v1.0",
        "生成日期：2026-05-07",
        "说明：当前未发现附件模板，本文件按通用课程提交格式生成。",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def split_table_row(line: str):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(line: str):
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def add_markdown_table(doc: Document, rows):
    parsed = [split_table_row(row) for row in rows if not is_separator(row)]
    if not parsed:
        return
    table = doc.add_table(rows=len(parsed), cols=max(len(row) for row in parsed))
    table.style = "Table Grid"
    for r_idx, row in enumerate(parsed):
        for c_idx, cell in enumerate(row):
            p = table.cell(r_idx, c_idx).paragraphs[0]
            run = p.add_run(cell)
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(9.5)
            if r_idx == 0:
                run.bold = True


def add_markdown(doc: Document, text: str):
    lines = text.splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.lstrip().startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_rows.append(lines[i])
                i += 1
            add_markdown_table(doc, table_rows)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip())
            p.paragraph_format.left_indent = Inches(0.25)
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(95, 105, 125)
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
        i += 1


def build_docx(output: str, title: str, subtitle: str, files):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    set_styles(doc)
    add_cover(doc, title, subtitle)
    for idx, filename in enumerate(files):
        if idx:
            doc.add_page_break()
        add_markdown(doc, (BASE / filename).read_text(encoding="utf-8"))
    doc.save(BASE / output)
    print(BASE / output)


def main():
    build_docx(
        "alethicode-technical-documents.docx",
        "Alethicode 技术文档合集",
        "项目计划、需求规格、概要设计、详细设计、测试计划与报告、项目总结",
        TECH_DOCS,
    )
    build_docx(
        "alethicode-practice-summary-report.docx",
        "Alethicode 实训总结报告",
        "课程实训总结提交版",
        ["07-practice-summary-report.md"],
    )
    build_docx(
        "alethicode-formal-review-documents.docx",
        "Alethicode 正式评审增强版技术文档",
        "六大软件工程技术文档、图表和正式评审附录",
        [
            "formal-review/alethicode-formal-review-all.md",
            "formal-review/08-truthfulness-audit.md",
        ],
    )


if __name__ == "__main__":
    main()
